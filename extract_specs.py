from docx import Document
import os

def extract():
    path1 = r'c:\Users\Innovatech\Desktop\PROJET____ing4GL____S8\MYprojet\Donnees_Registre_Cancer_Tlemcen.docx'
    path2 = r'c:\Users\Innovatech\Desktop\PROJET____ing4GL____S8\MYprojet\Scenarios_Registre_Cancer_Tlemcen.docx'
    
    with open('specs_full.txt', 'w', encoding='utf-8') as out:
        if os.path.exists(path1):
            doc1 = Document(path1)
            out.write("=== DATA SPECS ===\n")
            for p in doc1.paragraphs:
                out.write(p.text + "\n")
        
        if os.path.exists(path2):
            doc2 = Document(path2)
            out.write("\n\n=== SCENARIOS ===\n")
            for p in doc2.paragraphs:
                out.write(p.text + "\n")

if __name__ == "__main__":
    extract()
